# -*- coding: utf-8 -*-
"""Data visualization.ipynb

Automatically generated by Colaboratory.

Original file is located at
    https://colab.research.google.com/drive/1bONgR6JOM4x8Szblt3ZVBZXn7rYq5So7

# Basic loading
"""

# Commented out IPython magic to ensure Python compatibility.
import numpy as np # linear algebra
import pandas as pd # data processing, CSV file I/O (e.g. pd.read_csv)
import matplotlib.pyplot as plt # data visualization
import seaborn as sns # statistical data visualization
# %matplotlib inline

import warnings

warnings.filterwarnings('ignore')

# get data from drive
from google.colab import drive
drive.mount('/content/drive')

# Commented out IPython magic to ensure Python compatibility.
# %cd /content/drive/My Drive/Colab Notebooks/

df = pd.read_csv('Preprocessed_data.csv')
coffee_df=pd.read_csv('Coffee_data.csv')
rose_df=pd.read_csv('Rose_df.csv')
pomelo_df=pd.read_csv('Pomelo_df.csv')
ginger_df=pd.read_csv('Ginger_df.csv')
squash_df=pd.read_csv('Squash_df.csv')

df.info()
df.isnull().sum()

print(df.columns)

#df = df.astype({'Sales': int, 'Price': int, 'Rating record': int, 'Discount(%)': int})

import matplotlib.dates as mdates  # add import statement for DateFormatter

#change Sales to integer type
df['Sales'] = df['Sales'].astype(int)
#change Date to datetime type
from datetime import datetime
df['New_date']=pd.to_datetime(df['Date'], format='%Y/%m/%d',errors = 'coerce')
df = df.dropna(subset=['Date'])

df=df.drop(columns=['Date'])
df.columns = df.columns.str.replace('New_date', 'Date')

daily_sales = df.groupby('Date')['Sales'].sum().reset_index()
#daily_sales['Sales_rounded'] = round(daily_sales['Sales']/ 1000000, 2) * 1000
#daily_sales

"""# (Total sales) With value on each point """

# create a line plot of sales by day
plt.plot(daily_sales['Date'], daily_sales['Sales_rounded'], marker='o')

# show the sales value on each point of the line
for x, y in zip(daily_sales['Date'], daily_sales['Sales_rounded']):
    plt.text(x, y, y, ha='center', va='bottom', fontsize=8)

# set the x-axis format to display dates in yyyy/mm/dd format
date_format = mdates.DateFormatter('%d/%m')
plt.gca().xaxis.set_major_formatter(date_format)
plt.gcf().autofmt_xdate()

# add x and y axis labels and a title
plt.xlabel('Date')
plt.ylabel('Total sales')
plt.title('Total sales by Day')

# increase the size of the markers and line
plt.rcParams['lines.markersize'] = 6
plt.rcParams['lines.linewidth'] = 2


# set the figure size to 10 inches wide by 6 inches tall
plt.figure(figsize=(100, 6))

# show the plot
plt.show()

"""# (Total sales) Without value on each point"""

df

# create a line plot of sales by day
plt.plot(daily_sales['Date'], daily_sales['Sales'], marker='o',color='brown')

# set the x-axis format to display dates in yyyy/mm/dd format
date_format = mdates.DateFormatter('%d/%m')
plt.gca().xaxis.set_major_formatter(date_format)
plt.gcf().autofmt_xdate()

# add x and y axis labels and a title
plt.xlabel('Date')
plt.ylabel('Total sales')
plt.title('Total sales by Day')

# increase the size of the markers and line
plt.rcParams['lines.markersize'] = 3
plt.rcParams['lines.linewidth'] = 2


# set the figure size to 10 inches wide by 6 inches tall
plt.figure(figsize=(10, 6))

# save the plot as a PNG image
plt.savefig('Total_sales_by_day.png', dpi=600, bbox_inches='tight')

# show the plot
plt.show()

"""Save graph to doc"""

!pip install python-docx
import docx

plt.savefig('my_plot.png')

# Create a new Word document
doc = docx.Document()

# Add a new paragraph to the document
doc.add_paragraph('Here is my plot:')

# Add the plot image to the document
doc.add_picture('my_plot.png')

# Save the Word document
doc.save('my_doc.docx')

"""#Category sales"""

# Group the data by category and date, and calculate the daily sales
category_all = df.groupby(['Category', pd.Grouper(key='Date', freq='D')])['Sales'].sum().reset_index()

# Create a line graph with a legend note
fig, ax = plt.subplots(figsize=(10, 6))

for category in category_all['Category'].unique():
    data = category_all[category_all['Category'] == category]
    ax.plot(data['Date'], data['Sales'], marker='o', label=category)

ax.set_title('Daily Sales by Category')
ax.set_xlabel('Date')
ax.set_ylabel('Sales (items)')
ax.legend(title='Category', loc='upper left', bbox_to_anchor=(1.02, 1))
plt.rcParams['lines.markersize'] = 3

plt.show()

category_all

grouped_sales = df.groupby(['Category'])['Sales'].sum()
grouped_sales_df = grouped_sales.reset_index()
grouped_sales_df['Percentage']=(grouped_sales_df['Sales']/(grouped_sales_df['Sales'].sum()))*100

grouped_count = df.groupby(['Category'])['Sales'].count()
grouped_count_df = grouped_count.reset_index()
grouped_count_df['Percentage']=(grouped_count_df['Sales']/(grouped_count_df['Sales'].sum()))*100

grouped_count_df

grouped_sales_df

# create a pie chart of sales % using matplotlib
plt.pie(grouped_sales_df['Percentage'], labels=grouped_sales_df['Category'], autopct='%1.1f%%', startangle=90)

# add a title to the chart
plt.title('Sales by Category')

# show the chart
plt.show()

#create a pie chart of sales % using matplotlib
plt.pie(grouped_count_df['Percentage'], labels=grouped_count_df['Category'], autopct='%1.1f%%', startangle=90)

# add a title to the chart
plt.title('Percentage distribution of sales data collected by Category')

# show the chart
plt.show()

# Create a figure and axis object
fig, ax = plt.subplots()

# Create a bar graph
ax.bar(grouped_sales_df['Category'], grouped_sales_df['Sales'], color='brown')

# Set the title and axis labels
ax.set_title('Total sales by category')
ax.set_xlabel('Category')
ax.set_ylabel('Sales (items)')

# Add labels with values above the bars
for i, v in enumerate(grouped_sales_df['Sales']):
    ax.text(i, v+1, str(v), ha='center')

# Display the graph
plt.show()

# Create a figure and axis object
fig, ax = plt.subplots()

# Create a bar graph
ax.bar(grouped_count_df['Category'], grouped_count_df['Sales'], color='brown')

# Set the title and axis labels
ax.set_title('Number of Sales data collected by category')
ax.set_xlabel('Category')
ax.set_ylabel('Number of sales data collected')

# Add labels with values above the bars
for i, v in enumerate(grouped_count_df['Sales']):
    ax.text(i, v+1, str(v), ha='center')

# Display the graph
plt.show()

"""# Rating and Rating record"""

# group by Category and Rating, and count the number of occurrences
count_rating = df.groupby(['Category', 'Rating'])['Rating'].count()
count_rating_df = count_rating.reset_index()
# create a dictionary to map values to colors
color_map = {0.0:'orange', 4.0:'yellow',4.7: 'maroon', 4.9: 'lightgreen', 5.0: 'darkgreen'}
ascending=True
# pivot the Rating values into separate columns
count_rating_df = count_rating.unstack(level=-1)

# plot a stacked bar chart with the color of the Rating value
count_rating_df.plot(kind='bar', stacked=True, color=color_map)

# customize the plot
plt.xlabel('Category')
plt.ylabel('Count')
plt.title('Rating Distribution by Category')
plt.legend(title='Rating', loc='upper right', labels=color_map)
plt.show()

# Create a gradient color map
from matplotlib.colors import LinearSegmentedColormap
import matplotlib.dates as dates

cmap = LinearSegmentedColormap.from_list('mycmap', ['#070A52', '#F15A59'])

# Group the data by date and calculate the daily sales
daily_rate = df.groupby(pd.Grouper(key='Date', freq='D'))['Rating'].sum().reset_index()

# Calculate the difference in sales between consecutive days
daily_rate['Rate_diff'] = daily_rate['Rating'].diff()

# Create a bar chart of sales difference by day
fig, ax = plt.subplots()
ax.bar(daily_rate['Date'], daily_rate['Rate_diff'], color=cmap(daily_rate['Rate_diff']/daily_rate['Rate_diff'].max()))
ax.set_title('Rating Growth by Day')
ax.set_xlabel('Date')
ax.set_ylabel('Rating Difference')

# Set date format
date_format = '%d/%m'
ax.xaxis.set_major_formatter(dates.DateFormatter(date_format))


plt.show()

"""# Sales difference"""

# Create a gradient color map
from matplotlib.colors import LinearSegmentedColormap
import matplotlib.dates as dates

cmap = LinearSegmentedColormap.from_list('mycmap', ['#D21312', '#54B435'])

# Group the data by date and calculate the daily sales
daily_sales = df.groupby(pd.Grouper(key='Date', freq='D'))['Sales'].sum().reset_index()

# Calculate the difference in sales between consecutive days
daily_sales['Sales_diff'] = daily_sales['Sales'].diff()

# Create a bar chart of sales difference by day
fig, ax = plt.subplots()
ax.bar(daily_sales['Date'], daily_sales['Sales_diff'], color=cmap(daily_sales['Sales_diff']/daily_sales['Sales_diff'].max()))
ax.set_title('Sales Difference by Day')
ax.set_xlabel('Date')
ax.set_ylabel('Sales Difference')

# Set date format
date_format = '%d/%m'
ax.xaxis.set_major_formatter(dates.DateFormatter(date_format))

# Add line for average discount
# discount = df['Discount']
# ax.axhline(y=0, color='black', linestyle='-', linewidth=1)
# ax.axhline(y=discount * 100, color='red', linestyle='--', linewidth=1)
# ax.text(-0.5, discount * 100 + 5, 'Avg. Discount: {:.2f}%'.format(discount * 100), color='red', fontsize=10)

plt.show()

"""# Top lists """

ID_sales = df.groupby(['Product_ID','Category'])['Sales'].sum()
ID_sales_df = ID_sales.reset_index()
ID_sales_df

# Sort data by sales and get top 10 sellers
ID_sorted = ID_sales_df.sort_values('Sales', ascending=True)
ID_top10 = ID_sorted.tail(10)

# Create plot
fig, ax = plt.subplots()
ax.barh(ID_top10['Product_ID'], ID_top10['Sales'],color='orange')

# Set plot title and labels
ax.set_title('Top 10 Best-Seller')
ax.set_xlabel('Total sales (items)')
ax.set_ylabel('Product ID')

# Add data labels
for i, v in enumerate(ID_top10['Sales']):
    ax.text(v + 5 , i - 0.1, str(v), color='black', fontsize=8)

# Show plot
plt.show()

for category in category_all['Category'].unique():
    ID_top10 = category_all[category_all['Category'] == category]

# Sort data by sales and get top 10 sellers
ID_sorted = ID_sales_df.sort_values('Sales', ascending=True)
ID_top10 = ID_sorted.tail(10)
category_note= ID_top10['Category']

# Define colors for each category
colors = {'Coffee': 'brown', 'Pomelo': 'green', 'Squash': 'blue', 'Ginger': 'orange', 'Rose':'red'}

# Create plot
fig, ax = plt.subplots()
for index, row in ID_top10.iterrows():
    product_id = row['Product_ID']
    category = ID_sorted.loc[ID_sorted['Product_ID'] == product_id, 'Category'].iloc[0]
    color = colors[category]
    ax.barh(product_id, row['Sales'], color=color, label=category)

# Set plot title and labels
ax.set_title('Top 10 Best Seller')
ax.set_xlabel('Total sales (items)')
ax.set_ylabel('Product ID')
ax.set_title('Top 10 Best-Seller')
ax.legend(title='Category', loc='upper left', bbox_to_anchor=(1.02, 1))

# Add data labels
for i, v in enumerate(ID_top10['Sales']):
    ax.text(v + 5 , i - 0.1, str(v), color='black', fontsize=8)

# Show plot
plt.show()

ID_bot10 = ID_sorted.head(10)

# Create plot
fig, ax = plt.subplots()
ax.barh(ID_bot10['Product_ID'], ID_bot10['Sales'],color='maroon')

# Set plot title and labels
ax.set_title('Top 10 Worst-Seller')
ax.set_xlabel('Total sales (items)')
ax.set_ylabel('Product ID')

# Add data labels
for i, v in enumerate(ID_bot10['Sales']):
    ax.text(v + 3, i - 0.1, str(v), color='black', fontsize=8)

# Show plot
plt.show()

"""# Relationship plots (Scatter plot, Box plot, Heatmap)"""

sns.scatterplot(x='Sales', y='Discount(%)', data=df)
plt.show()
# to show the trend line of this relationship:
sns.regplot(x='Sales', y='Discount(%)', data=df)
sns.relplot

sns.scatterplot(x='Rating record', y='Discount(%)', data=df)
plt.show()
# to show the trend line of this relationship:
sns.regplot(x='Rating record', y='Discount(%)', data=df)
sns.relplot

sns.scatterplot(x='Price', y='Rating record', data=df)
plt.show()
# to show the trend line of this relationship:
sns.regplot(x='Price', y='Rating record', data=df)
sns.relplot

null_disc = df.dropna(subset=['Discount(%)'])

from scipy.stats import pearsonr
r, p = pearsonr(null_disc['Sales'], null_disc['Discount(%)'])

print("Pearson's r correlation coefficient:", r)
print("p-value:", p)

"""The Pearson's r correlation coefficient of -0.099 indicates a weak negative correlation between the two variables, "Sales" and "Discount(%)". A negative correlation means that as one variable increases, the other variable tends to decrease.

The p-value of 0.119 suggests that the correlation coefficient is not statistically significant at the standard 0.05 significance level, meaning there is a 11.9% chance that we would observe this correlation coefficient by chance alone. This indicates that we do not have sufficient evidence to conclude that there is a significant linear relationship between the two variables.

Overall, the result suggests that there is a weak negative relationship between "Sales" and "Discount(%)", but it is not strong enough to be statistically significant. It is important to note that correlation does not imply causation, and there may be other factors that are affecting the relationship between these variables. Further analysis and modeling may be necessary to understand the relationship more fully.
"""

df['Rating'].value_counts()

sns.boxplot(x=df['Rating'],y=df['Sales'])
plt.show()

plt.figure(figsize=(10,10)) # resize the figure size in plots
sns.heatmap(round(df.corr(),2),annot=True)
plt.show()

"""The code you provided is creating a heatmap using the Seaborn library to visualize the correlation matrix of a pandas DataFrame called df. Here's what you can learn from the plot:

The heatmap shows a color-coded matrix of the pairwise correlations between all the numeric columns in the DataFrame. Each cell in the matrix represents the correlation coefficient between two variables, with darker colors indicating stronger positive correlations and lighter colors indicating weaker correlations or negative correlations.

The diagonal cells of the matrix represent the correlation between a variable and itself, which is always equal to 1.

The annot=True parameter displays the correlation coefficients as text annotations in each cell of the matrix, making it easier to interpret the heatmap.

The round() function is used to round the correlation coefficients to two decimal places for readability.

By setting figsize=(10,10), the size of the heatmap figure is set to 10 by 10 inches.

Overall, the heatmap provides a visual summary of the pairwise correlations between all the variables in the DataFrame, allowing you to identify which variables are positively or negatively correlated with each other, and which variables may be less important in predicting the target variable. It can also help identify any multicollinearity issues between predictor variables that may affect the performance of a machine learning model.
"""